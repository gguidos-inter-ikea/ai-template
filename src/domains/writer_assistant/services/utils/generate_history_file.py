from docx import Document
from typing import List
from io import BytesIO

def generate_docx_file(chat_history: List[dict]):
    """
    Generates a .docx file from the chat history.

    Args:
        chat_history (List[dict]): The chat history to be included in the document.

    Returns:
        BytesIO: A stream containing the .docx file.
        str: The MIME type of the file.
    """
    doc = Document()
    doc.add_heading('Chat History', level=1)

    for message in chat_history:
        doc.add_heading("User:", level=2)
        doc.add_paragraph(message['prompt'])
        doc.add_heading("LLM:", level=2)
        doc.add_paragraph(message['generated_text'])
    
    file_stream= BytesIO()
    doc.save(file_stream)
    file_stream
    
    return file_stream, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'